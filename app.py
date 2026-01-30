"""
TruDiagnostic SOW Generator - Web Application
"""

import streamlit as st
from datetime import datetime, date
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

CONFIG = {
    "company_info": {"name": "Tru Diagnostic, Inc", "short_name": "TruDiagnostic"},
    "processing_types": {
        "epigenetic": {
            "name": "Epigenetic Processing",
            "sample_types": ["Blood Spot", "Whole Blood", "Buccal Swab"],
            "array_type": "MSA",
            "report_options": {
                "truage_truhealth": {"name": "Epigenetic + TruAge + TruHealth", "price": 300},
                "truage_only": {"name": "Epigenetic + TruAge Only", "price": 250},
                "truhealth_only": {"name": "Epigenetic + TruHealth Only", "price": 250}
            }
        },
        "genomic": {
            "name": "Genomic Processing",
            "sample_types": ["Blood Spot", "Whole Blood", "Buccal Swab"],
            "array_type": "GSAv3",
            "report_options": {
                "genomic_standard": {"name": "Genomic Processing (GSAv4ePgX)", "price": 100}
            }
        }
    },
    "report_only_options": {
        "truage_truhealth": {"name": "TruAge + TruHealth Report Only", "price": 95},
        "truhealth_only": {"name": "TruHealth Report Only", "price": 50},
        "truage_only": {"name": "TruAge Report Only", "price": 50}
    },
    "operational_services": {
        "kitting": {"name": "Kitting Services", "price": 5, "unit": "kit"},
        "3pl": {"name": "3PL/Fulfillment Services", "price": 15, "unit": "kit"},
        "customer_support": {"name": "Customer Support Services", "price": None, "unit": None, "included": True}
    },
    "bioinformatic_services": {
        "irb_tier1": {"name": "IRB Submission - Tier 1 (Partner drafts)", "price": 5000},
        "irb_tier2": {"name": "IRB Submission - Tier 2 (TruDiagnostic drafts)", "price": 10000},
        "publication_drafting": {"name": "Publication Drafting", "price": "Custom"},
        "publication_submission": {"name": "Publication Submission", "price": "Custom"},
        "interventional_trial": {"name": "Interventional Trial Analysis", "price": "Custom"},
        "algorithm_creation": {"name": "Custom Algorithm Development", "price": "Custom"},
        "algorithm_validation": {"name": "Algorithm Validation", "price": "Custom"}
    },
    "data_delivery_options": ["IDAT Files", "VCF Files", "CSV Files", "PDF Reports"]
}

st.set_page_config(page_title="TruDiagnostic SOW Generator", page_icon="🧬", layout="wide")

st.markdown("""
<style>
    .main-header {font-size: 2.5rem; font-weight: bold; color: #1E3A5F; margin-bottom: 0.5rem;}
    .sub-header {font-size: 1.2rem; color: #666; margin-bottom: 2rem;}
    .step-header {background-color: #f0f7ff; padding: 1rem; border-radius: 0.5rem; margin-bottom: 1rem; border-left: 4px solid #1E3A5F;}
    .summary-box {background-color: #f8f9fa; padding: 1.5rem; border-radius: 0.5rem; border: 1px solid #dee2e6;}
</style>
""", unsafe_allow_html=True)

def generate_sow_document(data):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("STATEMENT OF WORK")
    run.bold = True
    run.font.size = Pt(14)
    doc.add_paragraph()
    effective_date = data['effective_date'].strftime('%B %d, %Y') if data['effective_date'] else "[Effective Date]"
    intro = doc.add_paragraph()
    intro.add_run(f'This Statement of Work (this "SOW") is entered into as of {effective_date} ')
    intro.add_run(f'(the "Effective Date") by and between Tru Diagnostic, Inc ("Company") and ')
    intro.add_run(f'{data["partner_name"]} ("Partner"). This SOW is issued under and made part of the ')
    intro.add_run('Master Services Agreement (the "Agreement") between the Parties.')
    doc.add_paragraph("The specific Services to be provided, together with applicable fees, timelines, and other deal-specific details, are identified in Exhibit A (Services Engaged). Company shall have no obligation to perform any Service not expressly designated in Exhibit A.")
    p = doc.add_paragraph()
    p.add_run("1. Purpose. ").bold = True
    p.add_run("The purpose of this SOW is to define the scope of services that Company may provide to Partner, establish responsibilities of each Party, and set forth the commercial and operational framework under which the services will be delivered.")
    p = doc.add_paragraph()
    p.add_run("2. Term. ").bold = True
    p.add_run("This SOW will commence on the Effective Date and remain in effect until terminated in accordance with the MSA or this SOW.")
    p = doc.add_paragraph()
    p.add_run("3. Services Offered. ").bold = True
    p.add_run('Company offers the following categories of services. "Services" means the specific services selected by Partner and set forth in Exhibit A (Services Engaged).')
    section_num = 1
    if data['track'] == 'processing' and data['operational_services']:
        p = doc.add_paragraph()
        p.add_run(f"3.{section_num} Operational Services. ").bold = True
        p.add_run("If identified in Exhibit A, the Company shall provide operational services as described in Appendix 1.")
        section_num += 1
    if data['track'] == 'processing':
        p = doc.add_paragraph()
        p.add_run(f"3.{section_num} Laboratory Processing Services. ").bold = True
        p.add_run("If identified in Exhibit A, the Company shall provide the following services:")
        lab_services = ["Sample Receipt & Accessioning – Intake, verification, and LIMS entry of biological samples.", "Quality Control (QC) – Comprehensive assessment of DNA quality, quantity, and integrity prior to analysis.", "DNA Extraction & Quantification – Extraction from approved sample types with concentration/purity normalization."]
        if data['processing_type'] == 'epigenetic':
            lab_services.append("Epigenetic Processing – DNA methylation analysis using Array Manufacturer or Company-developed arrays.")
        else:
            lab_services.append("Genetic Processing – SNP and CNV detection using Array Manufacturer or custom genotyping arrays.")
        lab_services.append("Sample & Data Storage – Secure retention of residual materials and associated data.")
        for svc in lab_services:
            doc.add_paragraph(svc, style='List Bullet')
        section_num += 1
    p = doc.add_paragraph()
    p.add_run(f"3.{section_num} Bioinformatic Services. ").bold = True
    p.add_run("If identified in Exhibit A, the Company shall provide the following services:")
    bio_services = ["Data Processing – Conversion of raw array output into normalized and quality-controlled data sets.", "Algorithmic Analysis – Application of Company's proprietary algorithms.", "QC Reporting – Documentation of batch- and sample-level QC outputs."]
    if data['bioinformatic_services']:
        if any('irb' in svc for svc in data['bioinformatic_services']):
            bio_services.append("IRB Submission – Coordination of ethics review activities as specified in Exhibit A.")
        if 'publication_drafting' in data['bioinformatic_services']:
            bio_services.append("Publication Drafting – Manuscript preparation for peer-reviewed journals.")
        if 'publication_submission' in data['bioinformatic_services']:
            bio_services.append("Publication Submission – Journal submission support and coordination.")
        if 'interventional_trial' in data['bioinformatic_services']:
            bio_services.append("Interventional Trial Analysis – Statistical analysis of interventional study designs.")
        if 'algorithm_creation' in data['bioinformatic_services']:
            bio_services.append("Custom Algorithm Development – Development of tailored computational models.")
        if 'algorithm_validation' in data['bioinformatic_services']:
            bio_services.append("Algorithm Validation – Independent performance evaluation of algorithms.")
    for svc in bio_services:
        doc.add_paragraph(svc, style='List Bullet')
    section_num += 1
    p = doc.add_paragraph()
    p.add_run(f"3.{section_num} Reporting Services. ").bold = True
    p.add_run("If identified in Exhibit A, the Company shall provide the following services:")
    report_services = [f"Data Delivery – Secure transmission of data ({', '.join(data['data_delivery'])})."]
    if data['portal_access']:
        report_services.append("Technology Integration – Portal access and customized data delivery formats.")
    report_services.append("Data Security – Encryption, HIPAA compliance, and adherence to Company's technical safeguards.")
    for svc in report_services:
        doc.add_paragraph(svc, style='List Bullet')
    p = doc.add_paragraph()
    p.add_run("4. Dependencies. ").bold = True
    if data['track'] == 'processing':
        p.add_run("Partner acknowledges that certain services require timely provision of samples, manifests, metadata, and other inputs by Partner.")
    else:
        p.add_run("Partner acknowledges that Report Only services require timely provision of properly formatted .idat files and associated metadata.")
    p = doc.add_paragraph()
    p.add_run("5. Fees and Payment. ").bold = True
    p.add_run("Partner shall pay fees as set forth in Exhibit B.")
    p = doc.add_paragraph()
    p.add_run("6. Responsibilities.").bold = True
    doc.add_paragraph("Company Responsibilities:", style='List Bullet')
    for r in ["Perform services in accordance with validated SOPs and applicable regulations.", "Maintain regulatory certifications (e.g., CLIA) where required.", "Provide timely communication regarding QC failures, delays, or issues."]:
        p = doc.add_paragraph(r, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)
    doc.add_paragraph("Partner Responsibilities:", style='List Bullet')
    if data['track'] == 'processing':
        partner_resp = ["Ensure proper collection, de-identification, packaging, and shipping of samples.", "Retain all necessary patient consents and authorizations.", "Provide accurate manifests and metadata with each shipment."]
    else:
        partner_resp = ["Ensure proper formatting and quality of uploaded .idat files.", "Retain all necessary patient consents and authorizations.", "Provide accurate metadata files with each data submission."]
    for r in partner_resp:
        p = doc.add_paragraph(r, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)
    for num, (title, text) in enumerate([("Deliverables", "Deliverables may include raw and processed data files, QC reports, dashboards, or other outputs depending on services selected."), ("Confidentiality", "All Confidential Information exchanged will be handled in accordance with the MSA."), ("Intellectual Property", "Company retains ownership of its IP, methodologies, and algorithms. Partner retains ownership of samples and patient data."), ("Compliance", "Both Parties will comply with applicable laws including HIPAA, GDPR (if applicable), and FDA/CLIA/CAP requirements."), ("Termination", "This SOW may be terminated as provided in the MSA. Confidentiality and IP obligations survive termination."), ("Miscellaneous", "All other terms are governed by the MSA. This SOW may only be amended by written agreement.")], 7):
        p = doc.add_paragraph()
        p.add_run(f"{num}. {title}. ").bold = True
        p.add_run(text)
    doc.add_paragraph()
    doc.add_paragraph("IN WITNESS WHEREOF, the Parties have executed this Agreement as of the Effective Date.")
    doc.add_paragraph()
    table = doc.add_table(rows=5, cols=2)
    table.cell(0, 0).text = "TRU DIAGNOSTIC, INC."
    table.cell(0, 0).paragraphs[0].runs[0].bold = True
    table.cell(0, 1).text = data['partner_name'].upper()
    table.cell(0, 1).paragraphs[0].runs[0].bold = True
    for i, label in enumerate(["Signature: ________________", "Name: ________________", "Title: ________________", "Date: ________________"], 1):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = label
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("EXHIBIT A - SERVICES ENGAGED")
    run.bold = True
    run.font.size = Pt(14)
    doc.add_paragraph()
    doc.add_paragraph(f'Services engaged by {data["partner_name"]}:')
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = "Category"
    hdr[1].text = "Service"
    hdr[2].text = "Details"
    for cell in hdr:
        cell.paragraphs[0].runs[0].bold = True
    if data['track'] == 'processing':
        for op_key in data.get('operational_services', []):
            op = CONFIG['operational_services'].get(op_key, {})
            row = table.add_row().cells
            row[0].text = "Operational"
            row[1].text = f"[X] {op.get('name', op_key)}"
            row[2].text = f"${op.get('price', 'N/A')}/{op.get('unit', 'unit')}" if op.get('price') else "Included"
        proc = CONFIG['processing_types'].get(data['processing_type'], {})
        row = table.add_row().cells
        row[0].text = "Lab Processing"
        row[1].text = f"[X] {proc.get('name', 'Processing')}"
        row[2].text = f"Sample: {data.get('sample_type', 'N/A')}, Array: {proc.get('array_type', 'N/A')}"
    row = table.add_row().cells
    row[0].text = "Bioinformatics"
    row[1].text = "[X] Standard Bioinformatics"
    row[2].text = "Included"
    for bio_key in data.get('bioinformatic_services', []):
        bio = CONFIG['bioinformatic_services'].get(bio_key, {})
        row = table.add_row().cells
        row[0].text = ""
        row[1].text = f"[X] {bio.get('name', bio_key)}"
        price = bio.get('price', 'Custom')
        row[2].text = f"${price:,}" if isinstance(price, int) else str(price)
    row = table.add_row().cells
    row[0].text = "Reporting"
    row[1].text = f"[X] {data.get('report_name', 'Reports')}"
    row[2].text = f"${data.get('report_price', 0)}/sample"
    if data.get('portal_access'):
        row = table.add_row().cells
        row[0].text = ""
        row[1].text = "[X] Portal Access"
        row[2].text = "TruDiagnostic Provider Portal"
    if data['track'] == 'processing' and data.get('operational_services'):
        doc.add_page_break()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("APPENDIX 1 - SERVICE DESCRIPTIONS")
        run.bold = True
        run.font.size = Pt(14)
        doc.add_paragraph()
        if 'kitting' in data['operational_services']:
            p = doc.add_paragraph()
            p.add_run("Kitting Services").bold = True
            doc.add_paragraph("Custom Kitting Services, including design, assembly, and delivery of sample collection kits in accordance with specifications mutually agreed by the Parties.")
            for detail in ["Customized Kit Design - collaboration on packaging, branding, and labeling.", "Bill of Materials Documentation.", "Kit Identification - barcodes and Kit IDs.", "Quality Control Checks using statistical sampling.", "Kit Storage pending shipment."]:
                doc.add_paragraph(detail, style='List Bullet')
            doc.add_paragraph()
        if '3pl' in data['operational_services']:
            p = doc.add_paragraph()
            p.add_run("Third-Party Logistics (3PL) Services").bold = True
            doc.add_paragraph("Warehousing, order receipt, and order fulfillment services as designated by Partner.")
            for detail in ["Warehousing & Inventory Management. Partner remains sole owner of inventory.", "Order Receipt & Fulfillment through mutually agreed systems.", "Shipment tendering to Partner's preferred carriers."]:
                doc.add_paragraph(detail, style='List Bullet')
            doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("Laboratory Processing Services").bold = True
        doc.add_paragraph("Laboratory processing to enable generation of high-quality molecular data from Partner-provided samples.")
        lab_details = ["Sample Inspection and accession into LIMS within one (1) business day.", "Quality Control Checks including DNA concentration, purity ratios, and volume sufficiency.", "Reruns for Failed Analysis: One (1) rerun using residual sample material if initial analysis fails."]
        if data['processing_type'] == 'epigenetic':
            lab_details.extend(["DNA Extraction & Quantification from validated sample matrices.", "DNA Methylation Analysis using Illumina arrays with validated workflows."])
        else:
            lab_details.extend(["DNA Extraction & Normalization to array manufacturer specifications.", "Array-Based Genotyping using custom or commercial arrays."])
        for detail in lab_details:
            doc.add_paragraph(detail, style='List Bullet')
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("EXHIBIT B - PRICING AND PAYMENT TERMS")
    run.bold = True
    run.font.size = Pt(14)
    doc.add_paragraph()
    doc.add_paragraph("All fees are exclusive of applicable taxes, which shall be borne by Partner.")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("1. Service Fees").bold = True
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = "Service"
    hdr[1].text = "Unit"
    hdr[2].text = "Price (USD)"
    for cell in hdr:
        cell.paragraphs[0].runs[0].bold = True
    row = table.add_row().cells
    row[0].text = data.get('report_name', 'Report Services')
    row[1].text = "Per Sample"
    row[2].text = f"${data.get('report_price', 0)}"
    if data['track'] == 'processing':
        for op_key in data.get('operational_services', []):
            op = CONFIG['operational_services'].get(op_key, {})
            if op.get('price'):
                row = table.add_row().cells
                row[0].text = op.get('name', op_key)
                row[1].text = f"Per {op.get('unit', 'unit').title()}"
                row[2].text = f"${op['price']}"
    for bio_key in data.get('bioinformatic_services', []):
        bio = CONFIG['bioinformatic_services'].get(bio_key, {})
        row = table.add_row().cells
        row[0].text = bio.get('name', bio_key)
        row[1].text = "Flat Fee"
        price = bio.get('price', 'Custom')
        row[2].text = f"${price:,}" if isinstance(price, int) else "[To be determined]"
    doc.add_paragraph()
    if data.get('estimated_volume', 0) > 0:
        p = doc.add_paragraph()
        p.add_run("Estimated Volume: ").bold = True
        p.add_run(f"{data['estimated_volume']:,} samples")
        base_total = data.get('report_price', 0) * data['estimated_volume']
        p = doc.add_paragraph()
        p.add_run("Estimated Base Total: ").bold = True
        p.add_run(f"${base_total:,}")
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("2. Payment Terms").bold = True
    for term in ["Payment Terms: Net 30 days from invoice date.", "Invoicing: Upon completion of Services or monthly for ongoing Services.", "Late Payments: Past due amounts may accrue interest as specified in the Agreement.", "Taxes: All fees are exclusive of applicable taxes."]:
        doc.add_paragraph(term, style='List Bullet')
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("3. Delinquency Policy").bold = True
    doc.add_paragraph("Account Delinquent (1-30 Days): Lab processing restricted for samples linked to unpaid invoices.")
    doc.add_paragraph("Account Suspended (31+ Days): All lab processing blocked until balance resolved.")
    return doc

def main():
    st.markdown('<p class="main-header">🧬 TruDiagnostic SOW Generator</p>', unsafe_allow_html=True)
    st.markdown('<p class="sub-header">Generate customized Statements of Work for your clients</p>', unsafe_allow_html=True)
    if 'step' not in st.session_state:
        st.session_state.step = 1
    with st.sidebar:
        st.header("Progress")
        steps = ["Partner Info", "Service Type", "Service Details", "Add-on Services", "Review & Generate"]
        for i, step_name in enumerate(steps, 1):
            if i < st.session_state.step:
                st.success(f"✓ {step_name}")
            elif i == st.session_state.step:
                st.info(f"→ {step_name}")
            else:
                st.text(f"○ {step_name}")
    col1, col2 = st.columns([2, 1])
    with col1:
        if st.session_state.step == 1:
            st.markdown('<div class="step-header"><h3>Step 1: Partner Information</h3></div>', unsafe_allow_html=True)
            partner_name = st.text_input("Partner Legal Name *", placeholder="e.g., Acme Health Labs, Inc.")
            effective_date = st.date_input("Effective Date", value=date.today())
            st.markdown("---")
            if st.button("Next →", type="primary", disabled=not partner_name):
                st.session_state.partner_name = partner_name
                st.session_state.effective_date = effective_date
                st.session_state.step = 2
                st.rerun()
        elif st.session_state.step == 2:
            st.markdown('<div class="step-header"><h3>Step 2: Service Type</h3></div>', unsafe_allow_html=True)
            st.write("Will TruDiagnostic be processing biological samples, or will the partner upload pre-processed data?")
            track = st.radio("Select service type:", options=["processing", "report_only"], format_func=lambda x: "🧪 **Sample Processing** - We process samples and provide reports" if x == "processing" else "📊 **Report Only** - Partner uploads .idat files, we provide reports only", label_visibility="collapsed")
            st.markdown("---")
            col_a, col_b = st.columns(2)
            with col_a:
                if st.button("← Back"):
                    st.session_state.step = 1
                    st.rerun()
            with col_b:
                if st.button("Next →", type="primary"):
                    st.session_state.track = track
                    st.session_state.step = 3
                    st.rerun()
        elif st.session_state.step == 3:
            st.markdown('<div class="step-header"><h3>Step 3: Service Configuration</h3></div>', unsafe_allow_html=True)
            if st.session_state.track == 'processing':
                st.subheader("Processing Type")
                processing_type = st.radio("What type of processing?", options=["epigenetic", "genomic"], format_func=lambda x: f"🧬 {CONFIG['processing_types'][x]['name']}")
                st.subheader("Sample Type")
                sample_types = CONFIG['processing_types'][processing_type]['sample_types']
                sample_type = st.selectbox("Sample type:", sample_types)
                st.subheader("Report Package")
                report_options = CONFIG['processing_types'][processing_type]['report_options']
                report_choice = st.radio("Select report package:", options=list(report_options.keys()), format_func=lambda x: f"{report_options[x]['name']} - **${report_options[x]['price']}/sample**")
                report_name = report_options[report_choice]['name']
                report_price = report_options[report_choice]['price']
                st.subheader("Operational Services")
                operational_services = []
                col_ops1, col_ops2 = st.columns(2)
                with col_ops1:
                    if st.checkbox(f"Kitting Services (${CONFIG['operational_services']['kitting']['price']}/kit)"):
                        operational_services.append('kitting')
                with col_ops2:
                    if st.checkbox(f"3PL/Fulfillment (${CONFIG['operational_services']['3pl']['price']}/kit)"):
                        operational_services.append('3pl')
                if st.checkbox("Customer Support Services (Included)"):
                    operational_services.append('customer_support')
            else:
                st.subheader("Report Package")
                st.info("Partner will upload .idat files. Select the report package needed:")
                report_options = CONFIG['report_only_options']
                report_choice = st.radio("Select report package:", options=list(report_options.keys()), format_func=lambda x: f"{report_options[x]['name']} - **${report_options[x]['price']}/sample**")
                report_name = report_options[report_choice]['name']
                report_price = report_options[report_choice]['price']
                processing_type = None
                sample_type = None
                operational_services = []
            st.markdown("---")
            col_a, col_b = st.columns(2)
            with col_a:
                if st.button("← Back"):
                    st.session_state.step = 2
                    st.rerun()
            with col_b:
                if st.button("Next →", type="primary"):
                    st.session_state.processing_type = processing_type
                    st.session_state.sample_type = sample_type
                    st.session_state.report_choice = report_choice
                    st.session_state.report_name = report_name
                    st.session_state.report_price = report_price
                    st.session_state.operational_services = operational_services
                    st.session_state.step = 4
                    st.rerun()
        elif st.session_state.step == 4:
            st.markdown('<div class="step-header"><h3>Step 4: Additional Services</h3></div>', unsafe_allow_html=True)
            st.subheader("Bioinformatic Services")
            st.write("Standard bioinformatics (data processing, QC, algorithms) is included. Select any additional services:")
            bioinformatic_services = []
            col_bio1, col_bio2 = st.columns(2)
            with col_bio1:
                if st.checkbox("IRB Submission - Tier 1 ($5,000)"):
                    bioinformatic_services.append('irb_tier1')
                if st.checkbox("IRB Submission - Tier 2 ($10,000)"):
                    bioinformatic_services.append('irb_tier2')
                if st.checkbox("Publication Drafting (Custom)"):
                    bioinformatic_services.append('publication_drafting')
                if st.checkbox("Publication Submission (Custom)"):
                    bioinformatic_services.append('publication_submission')
            with col_bio2:
                if st.checkbox("Interventional Trial Analysis (Custom)"):
                    bioinformatic_services.append('interventional_trial')
                if st.checkbox("Custom Algorithm Development (Custom)"):
                    bioinformatic_services.append('algorithm_creation')
                if st.checkbox("Algorithm Validation (Custom)"):
                    bioinformatic_services.append('algorithm_validation')
            st.subheader("Data Delivery")
            data_delivery = st.multiselect("Select data formats needed:", options=CONFIG['data_delivery_options'], default=["PDF Reports"])
            portal_access = st.checkbox("TruDiagnostic Portal Access")
            st.subheader("Estimated Volume")
            estimated_volume = st.number_input("Estimated number of samples/reports:", min_value=0, step=100)
            st.markdown("---")
            col_a, col_b = st.columns(2)
            with col_a:
                if st.button("← Back"):
                    st.session_state.step = 3
                    st.rerun()
            with col_b:
                if st.button("Next →", type="primary"):
                    st.session_state.bioinformatic_services = bioinformatic_services
                    st.session_state.data_delivery = data_delivery if data_delivery else ["PDF Reports"]
                    st.session_state.portal_access = portal_access
                    st.session_state.estimated_volume = estimated_volume
                    st.session_state.step = 5
                    st.rerun()
        elif st.session_state.step == 5:
            st.markdown('<div class="step-header"><h3>Step 5: Review & Generate</h3></div>', unsafe_allow_html=True)
            st.markdown('<div class="summary-box">', unsafe_allow_html=True)
            st.subheader("📋 SOW Summary")
            col_sum1, col_sum2 = st.columns(2)
            with col_sum1:
                st.write(f"**Partner:** {st.session_state.partner_name}")
                st.write(f"**Effective Date:** {st.session_state.effective_date}")
                st.write(f"**Track:** {'Sample Processing' if st.session_state.track == 'processing' else 'Report Only'}")
                if st.session_state.track == 'processing':
                    proc_name = CONFIG['processing_types'].get(st.session_state.processing_type, {}).get('name', 'N/A')
                    st.write(f"**Processing:** {proc_name}")
                    st.write(f"**Sample Type:** {st.session_state.sample_type}")
                st.write(f"**Report Package:** {st.session_state.report_name}")
            with col_sum2:
                st.write(f"**Base Price:** ${st.session_state.report_price}/sample")
                if st.session_state.operational_services:
                    ops = [CONFIG['operational_services'][k]['name'] for k in st.session_state.operational_services]
                    st.write(f"**Operational:** {', '.join(ops)}")
                if st.session_state.bioinformatic_services:
                    bio = [CONFIG['bioinformatic_services'][k]['name'] for k in st.session_state.bioinformatic_services]
                    st.write(f"**Add-ons:** {', '.join(bio)}")
                if st.session_state.estimated_volume:
                    base_total = st.session_state.report_price * st.session_state.estimated_volume
                    st.write(f"**Volume:** {st.session_state.estimated_volume:,} samples")
                    st.write(f"**Est. Base Total:** ${base_total:,}")
            st.markdown('</div>', unsafe_allow_html=True)
            st.markdown("---")
            col_a, col_b, col_c = st.columns([1, 2, 1])
            with col_a:
                if st.button("← Back"):
                    st.session_state.step = 4
                    st.rerun()
            with col_b:
                if st.button("🔄 Generate SOW Document", type="primary", use_container_width=True):
                    with st.spinner("Generating document..."):
                        data = {'partner_name': st.session_state.partner_name, 'effective_date': st.session_state.effective_date, 'track': st.session_state.track, 'processing_type': st.session_state.processing_type, 'sample_type': st.session_state.sample_type, 'report_choice': st.session_state.report_choice, 'report_name': st.session_state.report_name, 'report_price': st.session_state.report_price, 'operational_services': st.session_state.operational_services, 'bioinformatic_services': st.session_state.bioinformatic_services, 'data_delivery': st.session_state.data_delivery, 'portal_access': st.session_state.portal_access, 'estimated_volume': st.session_state.estimated_volume}
                        doc = generate_sow_document(data)
                        buffer = BytesIO()
                        doc.save(buffer)
                        buffer.seek(0)
                        st.session_state.doc_buffer = buffer
                        st.session_state.doc_ready = True
                        st.rerun()
            if st.session_state.get('doc_ready'):
                st.success("✅ Document generated successfully!")
                filename = f"SOW_{st.session_state.partner_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"
                st.download_button(label="📥 Download SOW Document", data=st.session_state.doc_buffer, file_name=filename, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
                if st.button("🔄 Start New SOW"):
                    for key in list(st.session_state.keys()):
                        del st.session_state[key]
                    st.rerun()
    with col2:
        st.markdown("### 💡 Quick Tips")
        if st.session_state.step == 1:
            st.info("Enter the exact legal name as it should appear on the contract.")
        elif st.session_state.step == 2:
            st.info("**Sample Processing**: Client sends samples to TruDiagnostic.\n\n**Report Only**: Client uploads .idat files.")
        elif st.session_state.step == 3:
            if st.session_state.track == 'processing':
                st.info("**Epigenetic** = DNA methylation.\n\n**Genomic** = SNP/CNV detection.\n\nCannot combine in one SOW.")
            else:
                st.info("Report-only pricing is lower since no lab processing is required.")
        elif st.session_state.step == 4:
            st.info("Standard bioinformatics is always included. Additional services are optional add-ons.")
        elif st.session_state.step == 5:
            st.info("Review all details carefully before generating.")

if __name__ == "__main__":
    main()
